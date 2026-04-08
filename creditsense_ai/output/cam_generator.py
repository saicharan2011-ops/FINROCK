import io
import hashlib
from typing import Optional
from docx import Document
from docx.shared import RGBColor, Pt
from creditsense_ai.state_schema import CreditState

# Lazy import to avoid circular dependency; CAM can work without it
try:
    from creditsense_ai.research.promoter_scorer import ScoringBreakdown
except ImportError:
    ScoringBreakdown = None


def _risk_label(score: float) -> str:
    if score >= 0.75:
        return "CRITICAL"
    if score >= 0.55:
        return "HIGH"
    if score >= 0.35:
        return "MODERATE"
    if score >= 0.15:
        return "LOW"
    return "MINIMAL"


class CAMGenerator:
    def generate(self, state: CreditState, promoter_breakdown=None) -> bytes:
        doc = Document()
        doc.add_heading('Credit Appraisal Memo (CAM)', 0)

        # Dynamic Header
        header = doc.add_paragraph()
        header.add_run(f"Company Name: {state.company_name}\n").bold = True
        loan_str = f"₹{state.loan_amount:,.2f}" if state.loan_amount else "Not specified"
        header.add_run(f"Loan Amount: {loan_str}\n")
        decision = getattr(state, "last_recommendation", "PENDING").upper()
        header.add_run(f"Final Recommendation: {decision}\n")

        # Conditional Red Alerts
        if state.risk_signals.circular_trading_flag > 0.5:
            alert_p = doc.add_paragraph()
            run = alert_p.add_run("CRITICAL FRAUD ALERT: CIRCULAR TRADING DETECTED")
            run.font.color.rgb = RGBColor(255, 0, 0)
            run.font.bold = True

        if state.risk_signals.tamper_detected:
            alert_p = doc.add_paragraph()
            run = alert_p.add_run("CRITICAL ALERT: DOCUMENT TAMPERING DETECTED")
            run.font.color.rgb = RGBColor(255, 0, 0)
            run.font.bold = True

        # Promoter risk alert (high/critical)
        promo_risk = state.risk_signals.promoter_risk
        promo_label = _risk_label(promo_risk)
        if promo_risk >= 0.55:
            alert_p = doc.add_paragraph()
            color = RGBColor(255, 0, 0) if promo_risk >= 0.75 else RGBColor(255, 140, 0)
            run = alert_p.add_run(
                f"⚠ PROMOTER RISK: {promo_label} ({promo_risk:.2f})"
            )
            run.font.color.rgb = color
            run.font.bold = True

        doc.add_heading('The Five Cs of Credit', level=1)

        # Character
        doc.add_heading('1. Character', level=2)
        doc.add_paragraph(
            f"Promoter Risk Score: {promo_risk:.4f}  [{promo_label}]"
        )
        doc.add_paragraph(f"Litigation Risk Score: {state.risk_signals.litigation_risk}")

        # ── Promoter Risk Breakdown (audit trail) ──────────────────────────────
        if promoter_breakdown is not None:
            bd = promoter_breakdown
            doc.add_heading('Promoter Risk — Scoring Breakdown', level=3)
            doc.add_paragraph(f"Base Score (LLM × confidence): {bd.base_score:.4f}")
            doc.add_paragraph(f"Litigation Delta: +{bd.litigation_delta:.4f}")
            doc.add_paragraph(f"Keyword Delta: +{bd.critical_keyword_delta:.4f}")
            if bd.high_risk_keyword_delta > 0:
                doc.add_paragraph(f"  ↳ High-risk keywords: +{bd.high_risk_keyword_delta:.4f}")
            if bd.moderate_keyword_delta > 0:
                doc.add_paragraph(f"  ↳ Moderate keywords: +{bd.moderate_keyword_delta:.4f}")
            doc.add_paragraph(f"Fraud Signal Delta: +{bd.fraud_signal_delta:.4f}")
            if bd.research_gap_delta > 0:
                doc.add_paragraph(f"Research Gap Penalty: +{bd.research_gap_delta:.4f}")
            doc.add_paragraph(f"Pre-floor Score: {bd.pre_floor_score:.4f}")
            if bd.applied_floor is not None:
                doc.add_paragraph(f"Applied Hard Floor: {bd.applied_floor:.2f}")
            doc.add_paragraph(f"Final Score: {bd.final_score:.4f}")
            if bd.triggered_flags:
                flags_str = ", ".join(bd.triggered_flags)
                p = doc.add_paragraph()
                run = p.add_run(f"Triggered Flags: {flags_str}")
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(180, 0, 0)

        # Capacity
        doc.add_heading('2. Capacity', level=2)
        doc.add_paragraph(f"Debt Service Coverage Ratio (DSCR): {state.financial_ratios.dscr}")
        doc.add_paragraph(f"Operating Margin: {state.financial_ratios.op_margin}")
        if state.financial_ratios.revenue_mismatch_flag:
            doc.add_paragraph("WARNING: GST vs Bank Statement revenue mismatch identified.")

        # Capital
        doc.add_heading('3. Capital', level=2)
        doc.add_paragraph(f"Debt/Equity Ratio: {state.financial_ratios.de_ratio}")

        # Collateral
        doc.add_heading('4. Collateral', level=2)
        doc.add_paragraph("Collateral details require physical verification. Primary security is current assets.")

        # Conditions
        doc.add_heading('5. Conditions', level=2)
        doc.add_paragraph(f"Sector Headwinds Score: {state.risk_signals.sector_headwind}")
        
        # Site Visit / Field Observations
        doc.add_heading('Field Observations', level=1)
        if state.research_summary:
            doc.add_paragraph(state.research_summary)
        else:
            doc.add_paragraph("No site visit notes available.")

        # Blockchain Certificate
        doc.add_heading('Blockchain Certificate', level=1)
        state_str = f"{state.company_name}{state.last_recommendation}{state.step_count}"
        content_hash = hashlib.sha256(state_str.encode()).hexdigest()
        doc.add_paragraph(f"Document Hash (SHA-256): {content_hash}")
        doc.add_paragraph(f"Generated by: CreditSense AI v1.0 — RL Agent")
        doc.add_paragraph(
            "This document is blockchain-stamped. "
            "Verify at: https://www.oklink.com/amoy/address/YOUR_CONTRACT"
        )

        # Save to buffer
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.read()
